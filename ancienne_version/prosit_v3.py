import docx
from docx.shared import Inches
import os

tab_file = []
files = os.listdir('.')
for name in files:
    tab_file.append(name)

tab_file.remove('logo_cesi.jpg')
tab_file.remove('prosit_v3.py')
tab_file.remove('sauvegarde')
doc_name = tab_file[0]

print("transformation de : " + doc_name)

doc_save_number = str(input("numéro du prosit : "))

doc_original = docx.Document(doc_name)
size_doc = len(doc_original.paragraphs)

tab_categorie = ['Contexte','analyse du contexte ','Mots clés', 'Mots-clés','Problématique','Problématiques','Contraintes', 'Contrainte', 'Livrables', 'Généralisation',
	'Piste de solutions','Pistes de solution','Piste de solution','Pistes de solutions','Plans d’action','Plan d’actions','Plans d’actions', 'Plan d’action','Réalisation du plan d’action ']

doc_final = docx.Document()

doc_final.add_picture("logo_cesi.jpg", width=Inches(2.5))
a = doc_final.add_heading(doc_original.paragraphs[0].text, level=1)
a.alignment = 1

for i in range(1, size_doc):
	t = True
	for j in range(0, len(tab_categorie)):
		
		if doc_original.paragraphs[i].text.strip().lower() == tab_categorie[j].lower():  # cas ou la ligne est un titre de paragraphe
			doc_final.add_heading(tab_categorie[j], level=2)
			t = False

	if t == True:																		# cas ou c'est une ligne normale
		u = doc_final.add_paragraph(doc_original.paragraphs[i].text)
		u.paragraph_format.line_spacing = 1 # interligne à 0

doc_fini_name = "CER_prosit_" + doc_save_number + "_hugo_laplace.docx"
doc_final.save(doc_fini_name)

print("fichier editer avec succes")
os.system(doc_fini_name) # ouverture fichier fini

#https://stackoverflow.com/questions/24031011/python-docx-library-text-align
#https://python-docx.readthedocs.io/en/latest/