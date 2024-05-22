import docx
from docx.shared import Inches
import os
import openai
import re
import sys
from PyQt5.QtWidgets import QApplication, QFileDialog, QWidget
from PyQt5.QtCore import QDir

def get_docx_file():
    dialog.setFileMode(QFileDialog.AnyFile)
    dialog.setFilter(QDir.Files)

    if dialog.exec_():
        file_name = dialog.selectedFiles()
        if file_name[0].endswith('.docx'):
            print(file_name[0])
            return True, str(file_name[0])
        else:
            print("Erreur : Le fichier sélectionné n'est pas un docx")
            return False
    else:
        print("Erreur : Pas de fichier sélectionné")
        return False


def generate_text(prompt):
    
    openai.api_key = " "

    completions = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.5,
    )

    message = completions.choices[0].text
    return message.strip()

def find_part(string, pattern): # fonction pour verfifier si string existe dans pattern
    string_tot = r'[a-z ]+' + pattern
    match = re.search(string_tot, string)
    if match:
        return match.group()
    return None

def est_dans_tab_categorie(chaîne):
    tab_categorie = ['contexte','Mots clé', 'Mots-clé','Problématique','Contraintes', 'Livrables','Généralisation','Pistes de solutions','Piste de solutions', 'Plan d’action','Réalisation du plan d’action ']

    # Convertir la chaîne d'entrée en minuscules, normaliser au singulier et supprimer les points à la fin
    chaîne = chaîne.lower().rstrip('s').rstrip('.:')
    
    # Convertir les éléments du tableau en minuscules, normaliser au singulier et supprimer les points à la fin
    tab_categorie = [catégorie.lower().rstrip('s').rstrip('.:') for catégorie in tab_categorie]

    # Vérifier si la chaîne normalisée se trouve dans le tableau
    return chaîne in tab_categorie

####################################################################################

app = QApplication(sys.argv)
dialog = QFileDialog()  # Définissez la boîte de dialogue comme une variable globale
importOk, doc_name = get_docx_file()
dialog.setResult(QFileDialog.Accepted) # Fermez explicitement la boîte de dialogue en indiquant que l'opération a été acceptée.

print("Transformation de : " + doc_name)

# trouver le numéro du prosit
match = re.search(r'\d+', doc_name)
if match:
	var_num = match.group()
print("Prosit numéro : ", var_num)

doc_original = docx.Document(doc_name)
size_doc = len(doc_original.paragraphs)

#for paragraphe in doc_original.paragraphs:
   # print(paragraphe.text)


doc_final = docx.Document()

doc_final.add_picture("logo_cesi.jpg", width=Inches(2.5))
a = doc_final.add_heading(doc_original.paragraphs[0].text, level=1)
a.alignment = 1

var_cat_mc = False

for paragraphe in doc_original.paragraphs:
    if est_dans_tab_categorie(paragraphe.text) == True:
        print(paragraphe.text)

	
doc_fini_name = "CER_prosit_" + var_num + "_hugo_laplace.docx"
doc_final.save(doc_fini_name)

print("fichier editer avec succes")
os.system(doc_fini_name) # ouverture fichier fini

#https://stackoverflow.com/questions/24031011/python-docx-library-text-align
#https://python-docx.readthedocs.io/en/latest/
