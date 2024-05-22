import docx
from docx.shared import Inches
import os

doc_name = str(input("nom du document a traiter : "))
print(doc_name)

doc_original = docx.Document(doc_name)
size_doc = len(doc_original.paragraphs)

tab_categorie = ['Contexte','analyse du contexte ','Mots clés', 'Mots-clés','Problématique','Problématiques','Contraintes', 'Contrainte', 'Livrables', 'Généralisation',
	'Piste de solutions','Pistes de solution','Piste de solution','Pistes de solutions','Plans d’action','Plan d’actions','Plans d’actions', 'Plan d’action','Réalisation du plan d’action ']

doc_final = docx.Document()

doc_final.add_picture("logo_cesi.png", width=Inches(2.5))
a = doc_final.add_heading(doc_original.paragraphs[0].text, level=1)
a.alignment = 1

for i in range(1, size_doc):
	t = True
	for j in range(0, len(tab_categorie)):
		
		if doc_original.paragraphs[i].text.strip().lower() == tab_categorie[j].lower():
			doc_final.add_heading(tab_categorie[j], level=2)
			t = False

	if t == True:
		doc_final.add_paragraph(doc_original.paragraphs[i].text)

doc_final.save("EDIT_" + doc_name)
print("fichier editer avec succes")

#https://stackoverflow.com/questions/24031011/python-docx-library-text-align
#https://python-docx.readthedocs.io/en/latest/