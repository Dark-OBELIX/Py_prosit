import sys
from PyQt5.QtWidgets import QApplication, QFileDialog
from PyQt5.QtCore import QDir
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

class DocumentProcessor:
    def __init__(self):
        self.dialog = QFileDialog()
        self.first_line_processed = False
        self.process_problematic = False
        self.process_keywords = False

    def get_docx_file(self):
        # Open a file dialog to select a DOCX file
        self.dialog.setFileMode(QFileDialog.ExistingFiles)
        self.dialog.setFilter(QDir.Files)

        if self.dialog.exec_():
            file_name = self.dialog.selectedFiles()
            if file_name[0].endswith('.docx'):
                return True, str(file_name[0])
            else:
                print("Error: The selected file is not a DOCX")
                return False, ""
        else:
            print("Error: No file selected")
            return False, ""

    def find_words_in_docx(self, file_path, words):
        # Open the document
        doc = Document(file_path)

        # Convert words to lowercase for case-insensitive comparison
        words_lower = [word.lower() for word in words]

        # Initialize a list to store the results
        found_words = []

        # Initialize line number
        line_number = 0

        # Iterate through all paragraphs in the document
        for paragraph in doc.paragraphs:
            # Increment the line number
            line_number += 1
            # Convert paragraph text to lowercase
            paragraph_text_lower = paragraph.text.lower()
            # Check if any of the words are present in the paragraph text
            for word in words_lower:
                if word in paragraph_text_lower:
                    found_words.append((word, line_number))

        # Return the list of found words with their lines
        return found_words

    def copy_file_to_docx(self, source_file_path, titles):
        # Create a new Word document
        doc = Document()

        # Define the font style for the entire document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Open the source file
        source_doc = Document(source_file_path)

        # Iterate through each paragraph in the source document
        for index, paragraph in enumerate(source_doc.paragraphs):
            try:
                # Clean the text to remove incompatible characters
                clean_text = paragraph.text

                # Process the first line containing text
                if not self.first_line_processed and clean_text.strip():
                    doc.add_picture("logo_cesi.jpg")
                    p = doc.add_paragraph(clean_text)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.runs[0].bold = True
                    self.first_line_processed = True
                    continue

                # Check if the current paragraph should be a title
                title = next((t for t, l in titles if l == index + 1), None)
                if title:
                    # Capitalize the first letter of the title
                    title = self.capitalize_first_letter(title)
                    # Add the paragraph with the "Title 1" style
                    doc.add_heading(title, level=1)
                    self.process_problematic = (title.lower() == "problématique")
                    self.process_keywords = any(word in title.lower() for word in ["mots clés", "mots-clés", "mot-clés", "mots clefs"])
                elif self.process_problematic:
                    # Add paragraphs following the "Problématique" title in bold with 4 spaces before
                    p = doc.add_paragraph("    " + clean_text)
                    p.runs[0].bold = True
                elif self.process_keywords:
                    # Add paragraphs following the "Mots clés" title with 3 spaces and a hyphen
                    p = doc.add_paragraph("   - " + clean_text)
                else:
                    # Add the paragraph with the normal style
                    p = doc.add_paragraph(clean_text)
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
            except Exception as e:
                print(f"Error processing paragraph {index + 1}: {e}")
                continue

        # Save the Word document under the name "CER_prosit_<number>_hugo_laplace.docx"
        before, after = self.split_path(source_file_path)
        var_num = str(self.extract_number_from_string(after))
        doc_finished_name = "CER_prosit_" + var_num + "_hugo_laplace.docx"
        doc.save(before + "\\" + doc_finished_name)
        print("File saved as: ", before + "\\" + doc_finished_name)

    @staticmethod
    def split_path(file_path):
        last_slash_index = -1

        # Loop to find the index of the last '/' or '\'
        for i in range(len(file_path)):
            if file_path[i] == '/' or file_path[i] == '\\':
                last_slash_index = i

        # Extract what is before and after the last '/'
        if last_slash_index != -1:
            before_last_slash = file_path[:last_slash_index]
            after_last_slash = file_path[last_slash_index + 1:]
        else:
            before_last_slash = ""
            after_last_slash = file_path

        return before_last_slash, after_last_slash

    @staticmethod
    def extract_number_from_string(s):
        # Use a regular expression to find the number in the string
        match = re.search(r'\d+', s)
        if match:
            return int(match.group())
        else:
            return None

    @staticmethod
    def capitalize_first_letter(s):
        return s[:1].upper() + s[1:]

############################################################

app = QApplication(sys.argv)
processor = DocumentProcessor()

importOk, doc = processor.get_docx_file()
if importOk:
    words = ['Définition des mots-clefs', 'Mots clés', 'Mots-clés', 'Mot-clés', 'Mots clefs', 'Problématique', 'contexte', 'Contrainte', 'Livrable', 'Généralisation', 'Pistes de solutions', 'Plan d’action', 'Réalisation du plan d’action']
    found_words = processor.find_words_in_docx(doc, words)
    processor.copy_file_to_docx(doc, found_words)
